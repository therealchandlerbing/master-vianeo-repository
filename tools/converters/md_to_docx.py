#!/usr/bin/env python3
"""
VIANEO Markdown to DOCX Converter
=================================

Converts Markdown documents to professionally formatted DOCX files
following VIANEO styling guidelines.

Usage:
    python md_to_docx.py --input document.md --output document.docx
"""

import argparse
import re
from pathlib import Path
from typing import Dict, Any, Optional, List

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.constants import DocxStyles
from core.utils import clean_text


# =============================================================================
# CONVERTER CLASS
# =============================================================================

class MarkdownToDocxConverter:
    """Converts Markdown to professional DOCX format."""

    def __init__(self):
        self.styles = DocxStyles()
        self.doc = None
        self.current_list_level = 0

    def convert(self, markdown: str, output_path: Path) -> bool:
        """
        Convert markdown content to DOCX.

        Args:
            markdown: Markdown content
            output_path: Output DOCX file path

        Returns:
            True if successful
        """
        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed")
            return False

        self.doc = Document()
        self._setup_document()
        self._process_markdown(markdown)
        self.doc.save(str(output_path))
        return True

    def _setup_document(self) -> None:
        """Set up document margins and default styles."""
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _process_markdown(self, markdown: str) -> None:
        """Process markdown content line by line."""
        lines = markdown.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Headers
            if line.startswith('#'):
                self._add_heading(line)

            # Horizontal rule
            elif line.strip() in ['---', '***', '___']:
                self._add_horizontal_rule()

            # Table
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = []
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(table_lines)
                continue

            # Numbered list
            elif re.match(r'^\d+\.\s', line):
                self._add_numbered_list_item(line)

            # Bullet list
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                self._add_bullet_item(line)

            # Checkbox
            elif '[ ]' in line or '[x]' in line:
                self._add_checkbox_item(line)

            # Bold text line
            elif line.strip().startswith('**') and line.strip().endswith('**'):
                self._add_bold_paragraph(line)

            # Regular paragraph
            elif line.strip():
                self._add_paragraph(line)

            i += 1

    def _add_heading(self, line: str) -> None:
        """Add a heading based on # count."""
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if not match:
            return

        level = len(match.group(1))
        text = match.group(2)

        heading = self.doc.add_heading(text, level=min(level, 9))

        for run in heading.runs:
            run.font.name = self.styles.FONT_FAMILY
            run.font.color.rgb = RGBColor.from_string(self.styles.PRIMARY_BLUE)

            if level == 1:
                run.font.size = Pt(self.styles.HEADING1_SIZE)
            elif level == 2:
                run.font.size = Pt(self.styles.HEADING2_SIZE)
            else:
                run.font.size = Pt(12)

    def _add_horizontal_rule(self) -> None:
        """Add a horizontal rule."""
        para = self.doc.add_paragraph()
        para.add_run('_' * 60)
        for run in para.runs:
            run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

    def _add_table(self, lines: List[str]) -> None:
        """Add a table from markdown table lines."""
        if len(lines) < 2:
            return

        # Parse header
        headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
        if not headers:
            return

        # Skip separator line
        data_lines = lines[2:] if len(lines) > 2 else []

        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Add header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            # Shade header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), self.styles.TABLE_HEADER_BG)
            header_cells[i]._tc.get_or_add_tcPr().append(shading)

        # Add data rows
        for line in data_lines:
            cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells_data) == len(headers):
                row = table.add_row()
                for i, cell_text in enumerate(cells_data):
                    row.cells[i].text = self._parse_inline(cell_text)

    def _add_numbered_list_item(self, line: str) -> None:
        """Add a numbered list item."""
        match = re.match(r'^\d+\.\s+(.+)$', line)
        if match:
            para = self.doc.add_paragraph(style='List Number')
            self._add_inline_formatting(para, match.group(1))

    def _add_bullet_item(self, line: str) -> None:
        """Add a bullet list item."""
        text = re.sub(r'^[-*]\s+', '', line.strip())
        para = self.doc.add_paragraph(style='List Bullet')
        self._add_inline_formatting(para, text)

    def _add_checkbox_item(self, line: str) -> None:
        """Add a checkbox item."""
        checked = '[x]' in line.lower()
        text = re.sub(r'\[[ xX]\]\s*', '', line.strip())

        para = self.doc.add_paragraph()
        checkbox = '' if checked else ''
        para.add_run(f"{checkbox} ")
        self._add_inline_formatting(para, text)

    def _add_bold_paragraph(self, line: str) -> None:
        """Add a bold paragraph (label)."""
        text = line.strip().strip('*')
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(self.styles.BODY_SIZE)

    def _add_paragraph(self, line: str) -> None:
        """Add a regular paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        self._add_inline_formatting(para, line)

    def _add_inline_formatting(self, para, text: str) -> None:
        """Add text with inline formatting to paragraph."""
        # Pattern to match bold, italic, and code
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`\[]+)'

        for match in re.finditer(pattern, text):
            segment = match.group(0)

            if segment.startswith('**') and segment.endswith('**'):
                # Bold
                run = para.add_run(segment[2:-2])
                run.bold = True
            elif segment.startswith('*') and segment.endswith('*'):
                # Italic
                run = para.add_run(segment[1:-1])
                run.italic = True
            elif segment.startswith('`') and segment.endswith('`'):
                # Code
                run = para.add_run(segment[1:-1])
                run.font.name = 'Consolas'
            elif segment.startswith('['):
                # Link - extract text
                link_match = re.match(r'\[([^\]]+)\]', segment)
                if link_match:
                    run = para.add_run(link_match.group(1))
                    run.font.color.rgb = RGBColor.from_string("0066CC")
            else:
                # Regular text
                run = para.add_run(segment)

            run.font.name = self.styles.FONT_FAMILY
            run.font.size = Pt(self.styles.BODY_SIZE)
            if not hasattr(run, 'font') or run.font.color.rgb is None:
                run.font.color.rgb = RGBColor.from_string(self.styles.BODY_GRAY)

    def _parse_inline(self, text: str) -> str:
        """Strip markdown formatting for plain text."""
        # Remove bold
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove italic
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        # Remove code
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove links
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return clean_text(text)


# =============================================================================
# MAIN FUNCTION
# =============================================================================

def convert_md_to_docx(
    input_path: Path,
    output_path: Optional[Path] = None
) -> Optional[Path]:
    """
    Convert Markdown file to DOCX.

    Args:
        input_path: Path to input markdown file
        output_path: Path for output DOCX (default: same name with .docx)

    Returns:
        Output path if successful, None otherwise
    """
    input_path = Path(input_path)

    if output_path is None:
        output_path = input_path.with_suffix('.docx')

    # Read markdown
    with open(input_path, 'r', encoding='utf-8') as f:
        markdown = f.read()

    # Convert
    converter = MarkdownToDocxConverter()
    if converter.convert(markdown, output_path):
        print(f"Converted: {input_path} -> {output_path}")
        return output_path

    return None


# =============================================================================
# CLI
# =============================================================================

def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown to professional DOCX"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        required=True,
        help='Input Markdown file'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Output DOCX file'
    )

    args = parser.parse_args()

    result = convert_md_to_docx(args.input, args.output)
    if result:
        print(f"Success: {result}")
        return 0
    else:
        print("Conversion failed")
        return 1


if __name__ == '__main__':
    exit(main())
