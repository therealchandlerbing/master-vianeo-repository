#!/usr/bin/env python3
"""
VIANEO Diagnostic Comment Generator
====================================

Generates professional Step 10 VIANEO Diagnostic Comment documents
in both DOCX and Markdown formats.

Usage:
    python generate_diagnostic.py --input diagnostic_data.yaml --output DiagnosticComment
"""

import argparse
import json
import yaml
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.constants import DocxStyles, ScoreThresholds, VIANEO_DIMENSIONS
from core.utils import format_date, safe_filename, clean_text, load_data_file


# =============================================================================
# DATA MODELS
# =============================================================================

@dataclass
class DimensionScore:
    """Data for a single VIANEO dimension score."""
    name: str = ""
    score: float = 0.0
    interpretation: str = ""

    @property
    def status_keyword(self) -> str:
        return ScoreThresholds.get_status_keyword(self.score)

    @property
    def color_code(self) -> str:
        return ScoreThresholds.get_color_code(self.score)


@dataclass
class CriticalPathAction:
    """Data for a critical path action item."""
    action: str = ""
    owner: str = ""
    deadline: str = ""


@dataclass
class DiagnosticData:
    """Data structure for VIANEO Diagnostic Comment."""

    # Header
    project_name: str = ""
    date: str = ""
    overall_maturity: str = ""

    # Executive Diagnostic (6-8 sentences total)
    strengths: str = ""  # 1-2 sentences
    risks: str = ""      # 1-2 sentences
    near_term_actions: str = ""  # 2-3 sentences
    evidence_gaps: str = ""      # 1-2 sentences

    # Dimension Summary
    dimension_scores: List[DimensionScore] = field(default_factory=list)
    overall_status: str = ""

    # Critical Path Forward
    immediate_priorities: List[str] = field(default_factory=list)  # Weeks 1-4, max 3
    short_term_priorities: List[str] = field(default_factory=list)  # Months 2-3, max 4
    medium_term_priorities: List[str] = field(default_factory=list)  # Months 4-6, max 4
    success_metrics: List[str] = field(default_factory=list)  # max 6

    # Footer Metadata
    assessment_methodology: str = ""
    evidence_sources: str = ""
    next_review: str = ""


# =============================================================================
# DOCX GENERATION
# =============================================================================

class DiagnosticDocumentGenerator:
    """Generator for VIANEO Diagnostic Comment documents."""

    def __init__(self, data: DiagnosticData):
        self.data = data
        self.styles = DocxStyles()

    def generate_docx(self, output_path: Path) -> bool:
        """Generate professional DOCX diagnostic document."""
        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed")
            return False

        doc = Document()
        self._setup_document(doc)

        # Add content
        self._add_header(doc)
        self._add_executive_diagnostic(doc)
        self._add_dimension_summary(doc)
        self._add_critical_path(doc)
        self._add_footer_metadata(doc)

        # Save document
        doc.save(str(output_path))
        return True

    def _setup_document(self, doc: Document) -> None:
        """Set up document formatting."""
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_styled_heading(
        self,
        doc: Document,
        text: str,
        level: int = 1,
        color: str = None
    ) -> None:
        """Add a styled heading."""
        heading = doc.add_heading(text, level=level)

        color = color or self.styles.PRIMARY_BLUE
        for run in heading.runs:
            run.font.name = self.styles.FONT_FAMILY
            run.font.color.rgb = RGBColor.from_string(color)

    def _add_header(self, doc: Document) -> None:
        """Add document header."""
        # Title
        title = doc.add_heading(
            f"{self.data.project_name}: Vianeo Main Diagnostic Comment",
            level=0
        )
        for run in title.runs:
            run.font.name = self.styles.FONT_FAMILY
            run.font.size = Pt(self.styles.TITLE_SIZE)
            run.font.color.rgb = RGBColor.from_string(self.styles.PRIMARY_BLUE)

        # Metadata
        meta_items = [
            ("Date", self.data.date or format_date()),
            ("Assessment Framework", "Vianeo Business Model Evaluation Playbook"),
            ("Overall Maturity", self.data.overall_maturity)
        ]

        for label, value in meta_items:
            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(self.styles.METADATA_SIZE)
            run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

            run = para.add_run(value)
            run.font.size = Pt(self.styles.METADATA_SIZE)
            run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

        doc.add_paragraph()  # Spacer

    def _add_executive_diagnostic(self, doc: Document) -> None:
        """Add Executive Diagnostic section."""
        self._add_styled_heading(doc, "Executive Diagnostic", level=1)

        # Strengths
        self._add_styled_heading(doc, "Strengths", level=2, color=self.styles.MEDIUM_GRAY)
        para = doc.add_paragraph()
        run = para.add_run(clean_text(self.data.strengths))
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(self.styles.BODY_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.styles.BODY_GRAY)

        # Risks
        self._add_styled_heading(doc, "Risks", level=2, color=self.styles.MEDIUM_GRAY)
        para = doc.add_paragraph()
        run = para.add_run(clean_text(self.data.risks))
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(self.styles.BODY_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.styles.BODY_GRAY)

        # Near-term Actions
        self._add_styled_heading(
            doc, "Near-term Actions (30-60 days)",
            level=2, color=self.styles.MEDIUM_GRAY
        )
        para = doc.add_paragraph()
        run = para.add_run(clean_text(self.data.near_term_actions))
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(self.styles.BODY_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.styles.BODY_GRAY)

        # Evidence Gaps
        self._add_styled_heading(doc, "Evidence Gaps", level=2, color=self.styles.MEDIUM_GRAY)
        para = doc.add_paragraph()
        run = para.add_run(clean_text(self.data.evidence_gaps))
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(self.styles.BODY_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.styles.BODY_GRAY)

    def _add_dimension_summary(self, doc: Document) -> None:
        """Add Dimension Summary section with table."""
        self._add_styled_heading(doc, "Dimension Summary", level=1)

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Set column widths
        for i, width in enumerate([Inches(2), Inches(1), Inches(3.5)]):
            for cell in table.columns[i].cells:
                cell.width = width

        # Headers
        headers = ["Dimension", "Score", "Interpretation"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            # Add header shading
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), self.styles.TABLE_HEADER_BG)
            header_cells[i]._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for dim in self.data.dimension_scores:
            row = table.add_row()
            cells = row.cells

            # Dimension name
            cells[0].text = dim.name
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True

            # Score with color coding
            cells[1].text = f"{dim.score:.1f}/5"
            for para in cells[1].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
            # Add score shading
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), dim.color_code.replace('#', ''))
            cells[1]._tc.get_or_add_tcPr().append(shading)

            # Interpretation
            cells[2].text = f"{dim.status_keyword} - {dim.interpretation}"

        # Overall Status
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run("Overall Status: ")
        run.bold = True
        run.font.size = Pt(self.styles.BODY_SIZE)
        para.add_run(clean_text(self.data.overall_status))

    def _add_critical_path(self, doc: Document) -> None:
        """Add Critical Path Forward section."""
        self._add_styled_heading(doc, "Critical Path Forward", level=1)

        # Immediate Priority (Weeks 1-4)
        self._add_styled_heading(
            doc, "Immediate Priority (Weeks 1-4)",
            level=2, color=self.styles.MEDIUM_GRAY
        )
        for i, action in enumerate(self.data.immediate_priorities[:3], 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(action)

        # Short-term Priority (Months 2-3)
        self._add_styled_heading(
            doc, "Short-term Priority (Months 2-3)",
            level=2, color=self.styles.MEDIUM_GRAY
        )
        for i, action in enumerate(self.data.short_term_priorities[:4], 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(action)

        # Medium-term Priority (Months 4-6)
        self._add_styled_heading(
            doc, "Medium-term Priority (Months 4-6)",
            level=2, color=self.styles.MEDIUM_GRAY
        )
        for i, action in enumerate(self.data.medium_term_priorities[:4], 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(action)

        # Success Metrics
        self._add_styled_heading(
            doc, "Success Metrics",
            level=2, color=self.styles.MEDIUM_GRAY
        )
        for metric in self.data.success_metrics[:6]:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(metric)

    def _add_footer_metadata(self, doc: Document) -> None:
        """Add footer metadata section."""
        # Horizontal rule (simulated with paragraph border)
        doc.add_paragraph()

        # Metadata
        meta_items = [
            ("Assessment Methodology", self.data.assessment_methodology),
            ("Evidence Sources", self.data.evidence_sources),
            ("Next Review", self.data.next_review)
        ]

        for label, value in meta_items:
            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(self.styles.METADATA_SIZE)
            run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

            run = para.add_run(value)
            run.font.size = Pt(self.styles.METADATA_SIZE)
            run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)


# =============================================================================
# MARKDOWN GENERATION
# =============================================================================

def generate_markdown(data: DiagnosticData) -> str:
    """Generate markdown version of diagnostic comment."""
    md = f"""# {data.project_name}: Vianeo Main Diagnostic Comment

**Date:** {data.date or format_date()}
**Assessment Framework:** Vianeo Business Model Evaluation Playbook
**Overall Maturity:** {data.overall_maturity}

---

## Executive Diagnostic

### Strengths

{clean_text(data.strengths)}

### Risks

{clean_text(data.risks)}

### Near-term Actions (30-60 days)

{clean_text(data.near_term_actions)}

### Evidence Gaps

{clean_text(data.evidence_gaps)}

---

## Dimension Summary

| Dimension | Score | Interpretation |
|-----------|-------|----------------|
"""

    for dim in data.dimension_scores:
        status = ScoreThresholds.get_status_keyword(dim.score)
        md += f"| **{dim.name}** | {dim.score:.1f}/5 | {status} - {dim.interpretation} |\n"

    md += f"""
**Overall Status:** {clean_text(data.overall_status)}

---

## Critical Path Forward

### Immediate Priority (Weeks 1-4)

"""
    for i, action in enumerate(data.immediate_priorities[:3], 1):
        md += f"{i}. {action}\n"

    md += """
### Short-term Priority (Months 2-3)

"""
    for i, action in enumerate(data.short_term_priorities[:4], 1):
        md += f"{i}. {action}\n"

    md += """
### Medium-term Priority (Months 4-6)

"""
    for i, action in enumerate(data.medium_term_priorities[:4], 1):
        md += f"{i}. {action}\n"

    md += """
### Success Metrics

"""
    for metric in data.success_metrics[:6]:
        md += f"- {metric}\n"

    md += f"""
---

**Assessment Methodology:** {data.assessment_methodology}

**Evidence Sources:** {data.evidence_sources}

**Next Review:** {data.next_review}
"""

    return md


# =============================================================================
# MAIN FUNCTION
# =============================================================================

def generate_diagnostic(
    input_path: Optional[Path] = None,
    output_path: Optional[Path] = None,
    data: Optional[DiagnosticData] = None,
    output_format: str = "both"
) -> Dict[str, Path]:
    """
    Generate Diagnostic Comment document(s).

    Args:
        input_path: Path to input data file (JSON/YAML)
        output_path: Path for output file (without extension)
        data: DiagnosticData object (alternative to input_path)
        output_format: "docx", "md", or "both"

    Returns:
        Dict mapping format to output path
    """
    # Load data if not provided
    if data is None:
        if input_path is None:
            raise ValueError("Either input_path or data must be provided")

        raw_data = load_data_file(input_path)

        # Parse dimension scores
        dimension_scores = []
        for d in raw_data.get('dimension_scores', []):
            dimension_scores.append(DimensionScore(**d))

        data = DiagnosticData(
            project_name=raw_data.get('project_name', 'Project'),
            date=raw_data.get('date', ''),
            overall_maturity=raw_data.get('overall_maturity', ''),
            strengths=raw_data.get('strengths', ''),
            risks=raw_data.get('risks', ''),
            near_term_actions=raw_data.get('near_term_actions', ''),
            evidence_gaps=raw_data.get('evidence_gaps', ''),
            dimension_scores=dimension_scores,
            overall_status=raw_data.get('overall_status', ''),
            immediate_priorities=raw_data.get('immediate_priorities', []),
            short_term_priorities=raw_data.get('short_term_priorities', []),
            medium_term_priorities=raw_data.get('medium_term_priorities', []),
            success_metrics=raw_data.get('success_metrics', []),
            assessment_methodology=raw_data.get('assessment_methodology', ''),
            evidence_sources=raw_data.get('evidence_sources', ''),
            next_review=raw_data.get('next_review', '')
        )

    # Determine output path
    if output_path is None:
        project_name = safe_filename(data.project_name or "Project")
        output_path = Path(f"{project_name}_Vianeo_Diagnostic_Comment")

    output_path = Path(output_path)
    outputs = {}

    # Generate Markdown first (always available)
    if output_format in ["md", "both"]:
        md_path = output_path.with_suffix('.md')
        md_content = generate_markdown(data)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        outputs['md'] = md_path
        print(f"Generated Markdown: {md_path}")

    # Generate DOCX
    if output_format in ["docx", "both"] and DOCX_AVAILABLE:
        docx_path = output_path.with_suffix('.docx')
        generator = DiagnosticDocumentGenerator(data)
        if generator.generate_docx(docx_path):
            outputs['docx'] = docx_path
            print(f"Generated DOCX: {docx_path}")

    return outputs


# =============================================================================
# CLI
# =============================================================================

def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Generate VIANEO Diagnostic Comment documents"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        help='Input data file (JSON or YAML)'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Output file path (without extension)'
    )
    parser.add_argument(
        '--format', '-f',
        choices=['docx', 'md', 'both'],
        default='both',
        help='Output format (default: both)'
    )

    args = parser.parse_args()

    if args.input:
        outputs = generate_diagnostic(
            input_path=args.input,
            output_path=args.output,
            output_format=args.format
        )
        print(f"\nGenerated {len(outputs)} file(s)")
    else:
        print("No input file provided. Use --input to specify data file.")


if __name__ == '__main__':
    main()
