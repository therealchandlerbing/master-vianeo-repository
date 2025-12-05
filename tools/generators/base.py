"""
VIANEO Base Document Generator
==============================

Base class providing common document generation functionality
shared across all VIANEO document generators.
"""

from pathlib import Path
from typing import Optional, List, Any

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.constants import DocxStyles
from core.utils import clean_text

# Check for python-docx availability
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


def is_docx_available() -> bool:
    """Check if python-docx is available for document generation."""
    return DOCX_AVAILABLE


class BaseDocumentGenerator:
    """
    Base class for VIANEO document generators.

    Provides common document setup, styling, and utility methods
    used across all DOCX document generators.

    Subclasses should:
    1. Call super().__init__() in their __init__
    2. Override generate_docx() to add document-specific content
    3. Use the provided helper methods for consistent styling
    """

    def __init__(self):
        """Initialize base generator with default styles."""
        self.styles = DocxStyles()

    def _setup_document(self, doc: 'Document') -> None:
        """
        Set up document formatting with standard margins.

        Args:
            doc: The python-docx Document object
        """
        if not DOCX_AVAILABLE:
            return

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_styled_heading(
        self,
        doc: 'Document',
        text: str,
        level: int = 1,
        color: Optional[str] = None
    ) -> Any:
        """
        Add a styled heading to the document.

        Args:
            doc: The python-docx Document object
            text: Heading text
            level: Heading level (0-6, where 0 is title)
            color: Hex color string (without #), defaults to PRIMARY_BLUE

        Returns:
            The heading paragraph object
        """
        if not DOCX_AVAILABLE:
            return None

        heading = doc.add_heading(text, level=level)

        color = color or self.styles.PRIMARY_BLUE
        for run in heading.runs:
            run.font.name = self.styles.FONT_FAMILY
            run.font.color.rgb = RGBColor.from_string(color)

        return heading

    def _add_styled_paragraph(
        self,
        doc: 'Document',
        text: str,
        bold_label: Optional[str] = None,
        font_size: Optional[int] = None,
        color: Optional[str] = None
    ) -> Any:
        """
        Add a styled paragraph with optional bold label prefix.

        Args:
            doc: The python-docx Document object
            text: Paragraph text
            bold_label: Optional label to prefix (will be bolded)
            font_size: Font size in points (defaults to BODY_SIZE)
            color: Text color hex string (defaults to BODY_GRAY)

        Returns:
            The paragraph object
        """
        if not DOCX_AVAILABLE:
            return None

        para = doc.add_paragraph()
        font_size = font_size or self.styles.BODY_SIZE
        color = color or self.styles.BODY_GRAY

        if bold_label:
            run = para.add_run(f"{bold_label}: ")
            run.bold = True
            run.font.name = self.styles.FONT_FAMILY
            run.font.size = Pt(font_size)

        run = para.add_run(clean_text(text))
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(color)

        return para

    def _add_metadata_line(
        self,
        doc: 'Document',
        label: str,
        value: str
    ) -> Any:
        """
        Add a metadata line (label: value) with standard styling.

        Args:
            doc: The python-docx Document object
            label: The label text (will be bolded)
            value: The value text

        Returns:
            The paragraph object
        """
        if not DOCX_AVAILABLE:
            return None

        para = doc.add_paragraph()

        run = para.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(self.styles.METADATA_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

        run = para.add_run(value)
        run.font.size = Pt(self.styles.METADATA_SIZE)
        run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

        return para

    def _add_table_with_headers(
        self,
        doc: 'Document',
        headers: List[str],
        style: str = 'Table Grid'
    ) -> Any:
        """
        Add a table with styled headers.

        Args:
            doc: The python-docx Document object
            headers: List of column header strings
            style: Table style name (default: 'Table Grid')

        Returns:
            The table object
        """
        if not DOCX_AVAILABLE:
            return None

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = style

        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        return table

    def _add_shaded_header_cell(
        self,
        cell: Any,
        bg_color: Optional[str] = None
    ) -> None:
        """
        Add background shading to a table header cell.

        Args:
            cell: The table cell object
            bg_color: Background color hex string (without #)
        """
        if not DOCX_AVAILABLE:
            return

        bg_color = bg_color or self.styles.TABLE_HEADER_BG
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _add_checklist(
        self,
        doc: 'Document',
        items: List[str],
        heading_text: str = "Quality Checklist"
    ) -> None:
        """
        Add a quality checklist section.

        Args:
            doc: The python-docx Document object
            items: List of checklist item strings
            heading_text: Heading for the checklist section
        """
        if not DOCX_AVAILABLE:
            return

        self._add_styled_heading(doc, heading_text, level=4)
        for item in items:
            para = doc.add_paragraph()
            para.add_run(f"[ ] {item}")
            para.style = 'List Bullet'

    def _add_character_count(
        self,
        doc: 'Document',
        text: str,
        limit: int,
        field_name: str
    ) -> Any:
        """
        Add character count indicator.

        Args:
            doc: The python-docx Document object
            text: The text to count characters in
            limit: Character limit
            field_name: Name of the field for display

        Returns:
            The paragraph object
        """
        if not DOCX_AVAILABLE:
            return None

        count = len(text)
        status = "OK" if count <= limit else "OVER"

        para = doc.add_paragraph()
        run = para.add_run(f"[{field_name}: {count}/{limit} characters - {status}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)
        run.italic = True

        return para

    def generate_docx(self, output_path: Path) -> bool:
        """
        Generate DOCX document.

        Override this method in subclasses to implement
        document-specific content generation.

        Args:
            output_path: Path to save the DOCX file

        Returns:
            True if successful, False otherwise
        """
        raise NotImplementedError(
            "Subclasses must implement generate_docx()"
        )
