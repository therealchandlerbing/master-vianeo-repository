#!/usr/bin/env python3
"""
VIANEO Executive Brief Generator
================================

Generates professional Step 0 Executive Brief documents in DOCX format
from structured input data (JSON/YAML) or markdown templates.

Usage:
    python generate_executive_brief.py --input data.yaml --output Executive_Brief.docx
    python generate_executive_brief.py --input brief.md --output brief.docx
"""

import argparse
import json
import yaml
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.constants import CharacterLimits, DocxStyles, FileNaming
from core.utils import (
    validate_character_limit,
    format_date,
    safe_filename,
    clean_text,
    load_data_file,
    ValidationReport
)
from generators.base import BaseDocumentGenerator, is_docx_available, DOCX_AVAILABLE

# Import python-docx components if available
if DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
else:
    # Type hint stub when python-docx not available
    Document = None


# =============================================================================
# DATA MODELS
# =============================================================================

@dataclass
class ExecutiveBriefData:
    """Data structure for Executive Brief content."""

    # Project Information
    project_name: str = ""
    date_prepared: str = ""
    prepared_by: str = ""
    brief_id: str = ""

    # B1: Project Name and Tagline
    tagline: str = ""

    # B2: Problem Statement
    problem_description: str = ""
    problem_severity: int = 3
    problem_frequency: str = ""
    affected_population: str = ""

    # B3: Solution Overview
    solution_description: str = ""
    solution_type: str = "Product"
    key_features: List[str] = field(default_factory=list)

    # B4: Market and Users
    market_description: str = ""
    primary_users: str = ""
    primary_buyers: str = ""
    users_are_buyers: bool = True
    tam: str = ""
    sam: str = ""
    som: str = ""
    sizing_method: str = "Bottom-up"

    # B5: Business Model
    revenue_description: str = ""
    revenue_streams: List[str] = field(default_factory=list)
    cac: str = ""
    ltv: str = ""
    ltv_cac_ratio: str = ""

    # B6: Traction
    status_description: str = ""
    key_metrics: Dict[str, str] = field(default_factory=dict)
    customer_interviews: int = 0
    pilot_users: int = 0
    paying_customers: int = 0
    revenue_to_date: str = "$0"

    # B7: Team
    team_overview: str = ""
    team_members: List[Dict[str, str]] = field(default_factory=list)
    team_strengths: List[str] = field(default_factory=list)

    # Maturity Assessment
    maturity_stage: str = "IDEA"
    trl: int = 1

    # Evidence Log
    evidence_log: List[Dict[str, Any]] = field(default_factory=list)


# =============================================================================
# DOCX GENERATION
# =============================================================================

class ExecutiveBriefGenerator(BaseDocumentGenerator):
    """Generator for VIANEO Executive Brief documents."""

    def __init__(self, data: ExecutiveBriefData):
        super().__init__()
        self.data = data

    def generate_docx(self, output_path: Path) -> bool:
        """
        Generate professional DOCX Executive Brief.

        Args:
            output_path: Path to save DOCX file

        Returns:
            True if successful, False otherwise
        """
        if not is_docx_available():
            print("Error: python-docx not installed")
            return False

        doc = Document()

        # Set document properties
        self._setup_document(doc)

        # Add content
        self._add_header(doc)
        self._add_section_b1(doc)
        self._add_section_b2(doc)
        self._add_section_b3(doc)
        self._add_section_b4(doc)
        self._add_section_b5(doc)
        self._add_section_b6(doc)
        self._add_section_b7(doc)
        self._add_maturity_assessment(doc)
        self._add_evidence_log(doc)

        # Save document
        doc.save(str(output_path))
        return True

    # Note: _setup_document(), _add_styled_heading(), _add_styled_paragraph(),
    # _add_character_count(), and _add_checklist() are inherited from BaseDocumentGenerator

    def _add_header(self, doc: Document) -> None:
        """Add document header."""
        # Title
        title = doc.add_heading("Executive Brief Template", level=0)
        for run in title.runs:
            run.font.name = self.styles.FONT_FAMILY
            run.font.size = Pt(self.styles.TITLE_SIZE)
            run.font.color.rgb = RGBColor.from_string(self.styles.PRIMARY_BLUE)

        # Subtitle
        self._add_styled_heading(doc, "Project Information", level=2)

        # Metadata
        self._add_styled_paragraph(doc, self.data.project_name, "Project Name")
        self._add_styled_paragraph(doc, self.data.date_prepared or format_date(), "Date Prepared")
        self._add_styled_paragraph(doc, self.data.prepared_by, "Prepared By")
        self._add_styled_paragraph(doc, self.data.brief_id, "Executive Brief ID")

        doc.add_paragraph()  # Spacer

    def _add_section_b1(self, doc: Document) -> None:
        """Add Section B1: Project Name and Tagline."""
        self._add_styled_heading(doc, "Section B1: Project Name and Tagline", level=2)

        # Character limit note
        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 150 characters maximum (combined)")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.project_name, "Project Name")
        self._add_styled_paragraph(doc, self.data.tagline, "One-Line Tagline")

        combined = f"{self.data.project_name}: {self.data.tagline}"
        self._add_character_count(doc, combined, CharacterLimits.B1_PROJECT_NAME_TAGLINE, "B1")

        # Quality checklist
        self._add_checklist(doc, [
            "Name is clear and memorable",
            "Tagline is under 150 characters",
            "Tagline captures core value proposition"
        ])

    def _add_section_b2(self, doc: Document) -> None:
        """Add Section B2: Problem Statement."""
        self._add_styled_heading(doc, "Section B2: Problem Statement", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 300 characters maximum")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.problem_description, "Problem Description")
        self._add_character_count(
            doc, self.data.problem_description,
            CharacterLimits.B2_PROBLEM_STATEMENT, "B2"
        )

        # Problem characteristics
        self._add_styled_heading(doc, "Problem Characteristics", level=3)
        self._add_styled_paragraph(doc, f"{self.data.problem_severity}/5", "Severity")
        self._add_styled_paragraph(doc, self.data.problem_frequency, "Frequency")
        self._add_styled_paragraph(doc, self.data.affected_population, "Affected Population")

        self._add_checklist(doc, [
            "Problem is stated without reference to solution",
            "Problem statement is under 300 characters",
            "Severity and frequency are quantified",
            "Affected population is specified"
        ])

    def _add_section_b3(self, doc: Document) -> None:
        """Add Section B3: Solution Overview."""
        self._add_styled_heading(doc, "Section B3: Solution Overview", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 300 characters maximum")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.solution_description, "Solution Description")
        self._add_character_count(
            doc, self.data.solution_description,
            CharacterLimits.B3_SOLUTION_OVERVIEW, "B3"
        )

        self._add_styled_paragraph(doc, self.data.solution_type, "Solution Type")

        # Key features
        self._add_styled_heading(doc, "Key Features", level=3)
        for i, feature in enumerate(self.data.key_features[:6], 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(feature)

        self._add_checklist(doc, [
            "Solution description is under 300 characters",
            "Clear connection to problem statement",
            "Key features are listed",
            "Solution type is identified"
        ])

    def _add_section_b4(self, doc: Document) -> None:
        """Add Section B4: Market and Users."""
        self._add_styled_heading(doc, "Section B4: Market and Users", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 300 characters maximum")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.market_description, "Target Market")
        self._add_character_count(
            doc, self.data.market_description,
            CharacterLimits.B4_MARKET_DESCRIPTION, "B4"
        )

        # User/Buyer distinction
        self._add_styled_heading(doc, "User/Buyer Distinction", level=3)
        self._add_styled_paragraph(doc, self.data.primary_users, "Primary Users")
        self._add_styled_paragraph(doc, self.data.primary_buyers, "Primary Buyers")

        same_text = "Yes" if self.data.users_are_buyers else "No"
        self._add_styled_paragraph(doc, same_text, "Are users and buyers the same?")

        # Market sizing
        self._add_styled_heading(doc, "Market Sizing", level=3)
        self._add_styled_paragraph(doc, self.data.tam, "TAM (Total Addressable Market)")
        self._add_styled_paragraph(doc, self.data.sam, "SAM (Serviceable Addressable Market)")
        self._add_styled_paragraph(doc, self.data.som, "SOM (Serviceable Obtainable Market)")
        self._add_styled_paragraph(doc, self.data.sizing_method, "Sizing Method")

        self._add_checklist(doc, [
            "Market description is under 300 characters",
            "User/buyer distinction is clear",
            "Market sizes are quantified",
            "Sizing methodology is documented"
        ])

    def _add_section_b5(self, doc: Document) -> None:
        """Add Section B5: Business Model."""
        self._add_styled_heading(doc, "Section B5: Business Model", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 300 characters maximum")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.revenue_description, "Revenue Model")
        self._add_character_count(
            doc, self.data.revenue_description,
            CharacterLimits.B5_REVENUE_MODEL, "B5"
        )

        # Revenue streams
        self._add_styled_heading(doc, "Revenue Streams", level=3)
        for i, stream in enumerate(self.data.revenue_streams[:3], 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(stream)

        # Unit economics
        self._add_styled_heading(doc, "Unit Economics", level=3)
        self._add_styled_paragraph(doc, self.data.cac, "Customer Acquisition Cost (CAC)")
        self._add_styled_paragraph(doc, self.data.ltv, "Lifetime Value (LTV)")
        self._add_styled_paragraph(doc, self.data.ltv_cac_ratio, "LTV:CAC Ratio")

        self._add_checklist(doc, [
            "Revenue model is under 300 characters",
            "Revenue streams are listed",
            "Unit economics are quantified",
            "Business model is sustainable"
        ])

    def _add_section_b6(self, doc: Document) -> None:
        """Add Section B6: Traction and Validation."""
        self._add_styled_heading(doc, "Section B6: Traction and Validation", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 300 characters maximum")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.status_description, "Current Status")
        self._add_character_count(
            doc, self.data.status_description,
            CharacterLimits.B6_TRACTION_STATUS, "B6"
        )

        # Key metrics
        self._add_styled_heading(doc, "Key Metrics", level=3)
        for metric, value in self.data.key_metrics.items():
            self._add_styled_paragraph(doc, value, metric)

        # Validation evidence
        self._add_styled_heading(doc, "Validation Evidence", level=3)
        self._add_styled_paragraph(doc, str(self.data.customer_interviews), "Customer Interviews")
        self._add_styled_paragraph(doc, str(self.data.pilot_users), "Pilot Users")
        self._add_styled_paragraph(doc, str(self.data.paying_customers), "Paying Customers")
        self._add_styled_paragraph(doc, self.data.revenue_to_date, "Revenue to Date")

        self._add_checklist(doc, [
            "Status description is under 300 characters",
            "Key metrics are quantified",
            "Validation evidence is documented",
            "All claims are supported by numbers"
        ])

    def _add_section_b7(self, doc: Document) -> None:
        """Add Section B7: Team."""
        self._add_styled_heading(doc, "Section B7: Team", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Character Limit: 200 characters maximum")
        run.font.size = Pt(9)
        run.italic = True

        self._add_styled_paragraph(doc, self.data.team_overview, "Team Overview")
        self._add_character_count(
            doc, self.data.team_overview,
            CharacterLimits.B7_TEAM_OVERVIEW, "B7"
        )

        # Team members
        self._add_styled_heading(doc, "Key Team Members", level=3)
        for i, member in enumerate(self.data.team_members[:5], 1):
            name = member.get('name', 'Team Member')
            role = member.get('role', 'Role')
            experience = member.get('experience', 'Relevant experience')
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(f"{name}")
            run.bold = True
            para.add_run(f" - {role}: {experience}")

        # Team strengths
        self._add_styled_heading(doc, "Team Strengths", level=3)
        for strength in self.data.team_strengths[:5]:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(strength)

        self._add_checklist(doc, [
            "Team description is under 200 characters",
            "Key team members are listed with relevant experience",
            "Team strengths are identified",
            "Team composition is appropriate for project stage"
        ])

    def _add_maturity_assessment(self, doc: Document) -> None:
        """Add Maturity Assessment section."""
        self._add_styled_heading(doc, "Maturity Assessment", level=2)

        # Stage selection
        self._add_styled_heading(doc, "Current Stage", level=3)
        stages = ["IDEA", "PROTOTYPE", "PILOT", "EARLY_COMMERCIALIZATION", "GROWTH"]
        for stage in stages:
            check = "[x]" if stage == self.data.maturity_stage else "[ ]"
            para = doc.add_paragraph()
            para.add_run(f"{check} {stage.replace('_', ' ').title()}")

        # TRL
        self._add_styled_heading(doc, "Technology Readiness Level (TRL)", level=3)
        self._add_styled_paragraph(doc, f"{self.data.trl}/9", "Current TRL")

    def _add_evidence_log(self, doc: Document) -> None:
        """Add Evidence Log section."""
        self._add_styled_heading(doc, "Evidence Tracking", level=2)
        self._add_styled_heading(doc, "Evidence Log", level=3)

        if not self.data.evidence_log:
            para = doc.add_paragraph()
            para.add_run("No evidence entries logged yet.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Headers
        headers = ["Evidence ID", "Section", "Source Type", "Quality Rating", "Description", "Date"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        # Data rows
        for evidence in self.data.evidence_log:
            row_cells = table.add_row().cells
            row_cells[0].text = evidence.get('id', '')
            row_cells[1].text = evidence.get('section', '')
            row_cells[2].text = evidence.get('source_type', '')
            row_cells[3].text = str(evidence.get('quality_rating', ''))
            row_cells[4].text = evidence.get('description', '')
            row_cells[5].text = evidence.get('date', '')


# =============================================================================
# MARKDOWN GENERATION
# =============================================================================

def generate_markdown(data: ExecutiveBriefData) -> str:
    """Generate markdown version of Executive Brief."""
    combined_b1 = f"{data.project_name}: {data.tagline}"
    b1_count = len(combined_b1)

    md = f"""# Executive Brief Template

## Project Information
**Project Name:** {data.project_name}
**Date Prepared:** {data.date_prepared or format_date()}
**Prepared By:** {data.prepared_by}
**Executive Brief ID:** {data.brief_id}

---

## Section B1: Project Name and Tagline
**Character Limit:** 150 characters maximum

**Project Name:**
{data.project_name}

**One-Line Tagline:**
{data.tagline}

[Character count: {b1_count}/150]

**Quality Checklist:**
- [ ] Name is clear and memorable
- [ ] Tagline is under 150 characters
- [ ] Tagline captures core value proposition

---

## Section B2: Problem Statement
**Character Limit:** 300 characters maximum

**Problem Description:**
{data.problem_description}

[Character count: {len(data.problem_description)}/300]

**Problem Characteristics:**
- **Severity:** {data.problem_severity}/5
- **Frequency:** {data.problem_frequency}
- **Affected Population:** {data.affected_population}

**Quality Checklist:**
- [ ] Problem is stated without reference to solution
- [ ] Problem statement is under 300 characters
- [ ] Severity and frequency are quantified
- [ ] Affected population is specified

---

## Section B3: Solution Overview
**Character Limit:** 300 characters maximum

**Solution Description:**
{data.solution_description}

[Character count: {len(data.solution_description)}/300]

**Solution Type:**
- [x] {data.solution_type}

**Key Features:**
"""
    for i, feature in enumerate(data.key_features[:6], 1):
        md += f"{i}. {feature}\n"

    users_buyers = "Yes" if data.users_are_buyers else "No"

    md += f"""
**Quality Checklist:**
- [ ] Solution description is under 300 characters
- [ ] Clear connection to problem statement
- [ ] Key features are listed
- [ ] Solution type is identified

---

## Section B4: Market and Users
**Character Limit:** 300 characters maximum

**Target Market:**
{data.market_description}

[Character count: {len(data.market_description)}/300]

**User/Buyer Distinction:**
- **Primary Users:** {data.primary_users}
- **Primary Buyers:** {data.primary_buyers}
- **Are users and buyers the same?** {users_buyers}

**Market Sizing:**
- **TAM (Total Addressable Market):** {data.tam}
- **SAM (Serviceable Addressable Market):** {data.sam}
- **SOM (Serviceable Obtainable Market):** {data.som}

**Sizing Method:**
- [x] {data.sizing_method}

**Quality Checklist:**
- [ ] Market description is under 300 characters
- [ ] User/buyer distinction is clear
- [ ] Market sizes are quantified
- [ ] Sizing methodology is documented

---

## Section B5: Business Model
**Character Limit:** 300 characters maximum

**Revenue Model:**
{data.revenue_description}

[Character count: {len(data.revenue_description)}/300]

**Revenue Streams:**
"""
    for i, stream in enumerate(data.revenue_streams[:3], 1):
        md += f"{i}. {stream}\n"

    md += f"""
**Unit Economics:**
- **Customer Acquisition Cost (CAC):** {data.cac}
- **Lifetime Value (LTV):** {data.ltv}
- **LTV:CAC Ratio:** {data.ltv_cac_ratio}

**Quality Checklist:**
- [ ] Revenue model is under 300 characters
- [ ] Revenue streams are listed
- [ ] Unit economics are quantified
- [ ] Business model is sustainable

---

## Section B6: Traction and Validation
**Character Limit:** 300 characters maximum

**Current Status:**
{data.status_description}

[Character count: {len(data.status_description)}/300]

**Key Metrics:**
"""
    for metric, value in data.key_metrics.items():
        md += f"- **{metric}:** {value}\n"

    md += f"""
**Validation Evidence:**
- **Customer Interviews:** {data.customer_interviews}
- **Pilot Users:** {data.pilot_users}
- **Paying Customers:** {data.paying_customers}
- **Revenue to Date:** {data.revenue_to_date}

**Quality Checklist:**
- [ ] Status description is under 300 characters
- [ ] Key metrics are quantified
- [ ] Validation evidence is documented
- [ ] All claims are supported by numbers

---

## Section B7: Team
**Character Limit:** 200 characters maximum

**Team Overview:**
{data.team_overview}

[Character count: {len(data.team_overview)}/200]

**Key Team Members:**
"""
    for i, member in enumerate(data.team_members[:5], 1):
        name = member.get('name', 'Team Member')
        role = member.get('role', 'Role')
        experience = member.get('experience', 'Relevant experience')
        md += f"{i}. **{name}** - {role}: {experience}\n"

    md += """
**Team Strengths:**
"""
    for strength in data.team_strengths[:5]:
        md += f"- {strength}\n"

    md += f"""
**Quality Checklist:**
- [ ] Team description is under 200 characters
- [ ] Key team members are listed with relevant experience
- [ ] Team strengths are identified
- [ ] Team composition is appropriate for project stage

---

## Maturity Assessment

### Current Stage
"""
    stages = ["IDEA", "PROTOTYPE", "PILOT", "EARLY_COMMERCIALIZATION", "GROWTH"]
    for stage in stages:
        check = "[x]" if stage == data.maturity_stage else "[ ]"
        md += f"- {check} {stage.replace('_', ' ').title()}\n"

    md += f"""
### Technology Readiness Level (TRL)
**Current TRL:** {data.trl}

**TRL Scale:**
- TRL 1-3: Basic research and concept development
- TRL 4-6: Technology development and validation
- TRL 7-9: System demonstration and deployment

---

## Evidence Tracking

### Evidence Log
| Evidence ID | Section | Source Type | Quality Rating | Description | Date |
|-------------|---------|-------------|----------------|-------------|------|
"""
    for evidence in data.evidence_log:
        eid = evidence.get('id', '')
        section = evidence.get('section', '')
        source = evidence.get('source_type', '')
        quality = evidence.get('quality_rating', '')
        desc = evidence.get('description', '')
        date = evidence.get('date', '')
        md += f"| {eid} | {section} | {source} | {quality} | {desc} | {date} |\n"

    md += """
**Quality Rating Scale:**
- **5 - Gold Standard:** Published research, audited financials, signed contracts
- **4 - Strong:** Multiple customer testimonials, verified metrics
- **3 - Moderate:** Single-source testimonials, founder-reported metrics
- **2 - Weak:** Anecdotal evidence, unverified claims
- **1 - Very Weak:** No supporting evidence

---
"""

    return md


# =============================================================================
# MAIN FUNCTION
# =============================================================================

def generate_executive_brief(
    input_path: Optional[Path] = None,
    output_path: Optional[Path] = None,
    data: Optional[ExecutiveBriefData] = None,
    output_format: str = "both"
) -> Dict[str, Path]:
    """
    Generate Executive Brief document(s).

    Args:
        input_path: Path to input data file (JSON/YAML)
        output_path: Path for output file (without extension)
        data: ExecutiveBriefData object (alternative to input_path)
        output_format: "docx", "md", or "both"

    Returns:
        Dict mapping format to output path
    """
    # Load data if not provided
    if data is None:
        if input_path is None:
            raise ValueError("Either input_path or data must be provided")

        raw_data = load_data_file(input_path)
        data = ExecutiveBriefData(**raw_data)

    # Determine output path
    if output_path is None:
        project_name = safe_filename(data.project_name or "Project")
        date_str = format_date()
        output_path = Path(f"Executive_Brief_{project_name}_{date_str}")

    output_path = Path(output_path)
    outputs = {}

    # Generate DOCX
    if output_format in ["docx", "both"] and is_docx_available():
        docx_path = output_path.with_suffix('.docx')
        generator = ExecutiveBriefGenerator(data)
        if generator.generate_docx(docx_path):
            outputs['docx'] = docx_path
            print(f"Generated DOCX: {docx_path}")

    # Generate Markdown
    if output_format in ["md", "both"]:
        md_path = output_path.with_suffix('.md')
        md_content = generate_markdown(data)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        outputs['md'] = md_path
        print(f"Generated Markdown: {md_path}")

    return outputs


# =============================================================================
# CLI
# =============================================================================

def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Generate VIANEO Executive Brief documents"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        help='Input data file (JSON or YAML)'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Output file path (without extension)'
    )
    parser.add_argument(
        '--format', '-f',
        choices=['docx', 'md', 'both'],
        default='both',
        help='Output format (default: both)'
    )

    args = parser.parse_args()

    if args.input:
        outputs = generate_executive_brief(
            input_path=args.input,
            output_path=args.output,
            output_format=args.format
        )
        print(f"\nGenerated {len(outputs)} file(s)")
    else:
        print("No input file provided. Use --input to specify data file.")
        print("\nExample:")
        print("  python generate_executive_brief.py --input project_data.yaml --output brief")


if __name__ == '__main__':
    main()
