#!/usr/bin/env python3
"""
VIANEO Persona Document Generator
=================================

Generates professional Step 6 Persona Development documents in DOCX format
from structured input data.

Usage:
    python generate_personas.py --input personas.yaml --output CompanyName_Personas.docx
"""

import argparse
import json
import yaml
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field
from enum import Enum

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.constants import CharacterLimits, DocxStyles
from core.utils import format_date, safe_filename, clean_text, count_characters


# =============================================================================
# DATA MODELS
# =============================================================================

class ValidationStatus(Enum):
    """Persona validation status based on interview count."""
    VALIDATED = "VALIDATED"           # 5+ interviews
    PARTIALLY_VALIDATED = "PARTIALLY_VALIDATED"  # 1-4 interviews
    NEEDS_VALIDATION = "NEEDS_VALIDATION"        # 0 interviews

    @classmethod
    def from_interview_count(cls, count: int) -> 'ValidationStatus':
        if count >= 5:
            return cls.VALIDATED
        elif count >= 1:
            return cls.PARTIALLY_VALIDATED
        else:
            return cls.NEEDS_VALIDATION


@dataclass
class PersonaData:
    """Data structure for a single persona."""

    # Identity
    first_name: str = ""
    age: int = 35

    # Profile
    life_motivations: str = ""
    personality_values: str = ""

    # Field of Application
    thinks_feels: str = ""
    observes: str = ""
    does: str = ""
    others_say: str = ""

    # Activities & Challenges (60 char limit each)
    tasks: List[str] = field(default_factory=list)
    pains: List[str] = field(default_factory=list)
    expectations: List[str] = field(default_factory=list)

    # Current Solutions
    current_solutions_description: str = ""
    current_tools: List[Dict[str, str]] = field(default_factory=list)

    # Validation
    interview_count: int = 0
    interview_date_range: str = ""

    @property
    def validation_status(self) -> ValidationStatus:
        return ValidationStatus.from_interview_count(self.interview_count)


@dataclass
class PersonaDocumentData:
    """Data structure for full persona document."""

    # Document info
    company_name: str = ""
    project_subtitle: str = ""
    prepared_date: str = ""

    # Research overview
    research_overview: str = ""
    critical_gaps: str = ""

    # Personas (3-5 required)
    personas: List[PersonaData] = field(default_factory=list)


# =============================================================================
# STYLING HELPERS
# =============================================================================

def get_validation_style(status: ValidationStatus) -> Dict[str, str]:
    """Get styling for validation status badge."""
    styles = {
        ValidationStatus.VALIDATED: {
            "bg_color": "E8F4EA",
            "text_color": "2D7A3E",
            "icon": "VALIDATED",
            "text": "VALIDATED"
        },
        ValidationStatus.PARTIALLY_VALIDATED: {
            "bg_color": "FFF4E6",
            "text_color": "CC7A00",
            "icon": "PARTIAL",
            "text": "PARTIALLY VALIDATED"
        },
        ValidationStatus.NEEDS_VALIDATION: {
            "bg_color": "FFEBEE",
            "text_color": "C62828",
            "icon": "NEEDS",
            "text": "NEEDS VALIDATION"
        }
    }
    return styles[status]


# =============================================================================
# DOCX GENERATION
# =============================================================================

class PersonaDocumentGenerator:
    """Generator for VIANEO Persona documents."""

    def __init__(self, data: PersonaDocumentData):
        self.data = data
        self.styles = DocxStyles()

    def generate_docx(self, output_path: Path) -> bool:
        """Generate professional DOCX persona document."""
        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed")
            return False

        doc = Document()
        self._setup_document(doc)

        # Add content pages
        self._add_cover_page(doc)
        doc.add_page_break()
        self._add_evidence_base_page(doc)

        # Add persona pages
        for i, persona in enumerate(self.data.personas, 1):
            doc.add_page_break()
            self._add_persona_page(doc, persona, i)

        # Save document
        doc.save(str(output_path))
        return True

    def _setup_document(self, doc: Document) -> None:
        """Set up document formatting."""
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_cover_page(self, doc: Document) -> None:
        """Add cover page."""
        # Company name
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(self.data.company_name)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("2E5C8A")
        run.font.name = self.styles.FONT_FAMILY

        # Subtitle
        para = doc.add_paragraph()
        run = para.add_run("Vianeo Business Model Evaluation")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string("6C757D")

        # Add spacer
        doc.add_paragraph()
        doc.add_paragraph()

        # Title
        para = doc.add_paragraph()
        run = para.add_run("USER PERSONA ANALYSIS")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("1a365d")

        # Project subtitle
        if self.data.project_subtitle:
            para = doc.add_paragraph()
            run = para.add_run(self.data.project_subtitle)
            run.font.size = Pt(14)
            run.font.italic = True
            run.font.color.rgb = RGBColor.from_string("4a5568")

        # Add spacer
        for _ in range(10):
            doc.add_paragraph()

        # Metadata
        para = doc.add_paragraph()
        run = para.add_run(f"Prepared: {self.data.prepared_date or format_date(format_str='%B %Y')}")
        run.font.size = Pt(11)
        run.bold = True

        total_interviews = sum(p.interview_count for p in self.data.personas)
        para = doc.add_paragraph()
        run = para.add_run(f"Validation Status: {len(self.data.personas)} personas based on {total_interviews} total interviews")
        run.font.size = Pt(11)

        para = doc.add_paragraph()
        run = para.add_run("Evaluation Phase: Desirability Assessment")
        run.font.size = Pt(11)

        doc.add_paragraph()

        para = doc.add_paragraph()
        run = para.add_run("VIANEO Framework v2.0")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("6C757D")

    def _add_evidence_base_page(self, doc: Document) -> None:
        """Add evidence base summary page."""
        # Heading
        heading = doc.add_heading("Evidence Base Summary", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(self.styles.PRIMARY_BLUE)

        # Research overview
        doc.add_heading("Research Overview", level=2)
        para = doc.add_paragraph()
        para.add_run(self.data.research_overview)

        # Validation breakdown table
        doc.add_heading("Validation Breakdown", level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Persona", "Interview Count", "Date Range", "Validation Status"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        # Data rows
        for persona in self.data.personas:
            row = table.add_row()
            row.cells[0].text = persona.first_name
            row.cells[1].text = str(persona.interview_count)
            row.cells[2].text = persona.interview_date_range
            status_style = get_validation_style(persona.validation_status)
            row.cells[3].text = status_style["text"]

        # Critical gaps
        doc.add_heading("Critical Gaps & Next Steps", level=2)
        para = doc.add_paragraph()
        para.add_run(self.data.critical_gaps)

    def _add_persona_page(self, doc: Document, persona: PersonaData, number: int) -> None:
        """Add individual persona page."""
        # Header with validation badge
        heading = doc.add_heading(f"PERSONA {number}: {persona.first_name.upper()}", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(self.styles.PRIMARY_BLUE)

        # Validation badge
        status_style = get_validation_style(persona.validation_status)
        para = doc.add_paragraph()
        run = para.add_run(f"[{status_style['icon']}] {status_style['text']}")
        run.font.bold = True
        run.font.size = Pt(11)

        para = doc.add_paragraph()
        if persona.interview_count > 0:
            run = para.add_run(f"Based on {persona.interview_count} interviews ({persona.interview_date_range})")
        else:
            run = para.add_run("No interviews conducted yet")
        run.font.size = Pt(10)
        run.font.italic = True

        # Requester Profile
        doc.add_heading("Requester Profile", level=2)

        para = doc.add_paragraph()
        run = para.add_run("First Name: ")
        run.bold = True
        para.add_run(persona.first_name)

        para = doc.add_paragraph()
        run = para.add_run("Age: ")
        run.bold = True
        para.add_run(str(persona.age))

        para = doc.add_paragraph()
        run = para.add_run("Life & Motivations: ")
        run.bold = True
        doc.add_paragraph(persona.life_motivations)

        para = doc.add_paragraph()
        run = para.add_run("Personality & Values: ")
        run.bold = True
        doc.add_paragraph(persona.personality_values)

        # Field of Application
        doc.add_heading("Field of Application", level=2)

        sections = [
            ("Thinks & Feels:", persona.thinks_feels),
            ("Observes:", persona.observes),
            ("Does:", persona.does),
            ("Others Say:", persona.others_say)
        ]

        for label, content in sections:
            para = doc.add_paragraph()
            run = para.add_run(label)
            run.bold = True
            doc.add_paragraph(content)

        # Activities & Challenges
        doc.add_heading("Activities & Challenges", level=2)

        # Tasks
        para = doc.add_paragraph()
        run = para.add_run("Tasks (Jobs to be done):")
        run.bold = True
        run.font.size = Pt(11)

        for task in persona.tasks[:6]:
            para = doc.add_paragraph(style='List Bullet')
            # Validate character limit
            if len(task) > 60:
                task = task[:57] + "..."
            para.add_run(task)

        # Pains
        para = doc.add_paragraph()
        run = para.add_run("Pains (Problems to eliminate):")
        run.bold = True
        run.font.size = Pt(11)

        for pain in persona.pains[:6]:
            para = doc.add_paragraph(style='List Bullet')
            if len(pain) > 60:
                pain = pain[:57] + "..."
            para.add_run(pain)

        # Expectations
        para = doc.add_paragraph()
        run = para.add_run("Expectations (Desired outcomes):")
        run.bold = True
        run.font.size = Pt(11)

        for exp in persona.expectations[:4]:
            para = doc.add_paragraph(style='List Bullet')
            if len(exp) > 60:
                exp = exp[:57] + "..."
            para.add_run(exp)

        # Current Solutions
        doc.add_heading("Current Solutions", level=2)
        doc.add_paragraph(persona.current_solutions_description)

        if persona.current_tools:
            para = doc.add_paragraph()
            run = para.add_run("Current Tools:")
            run.bold = True

            for tool in persona.current_tools:
                para = doc.add_paragraph(style='List Bullet')
                name = tool.get('name', 'Tool')
                limitation = tool.get('limitation', 'Limitation')
                run = para.add_run(f"{name}: ")
                run.bold = True
                para.add_run(limitation)


# =============================================================================
# MARKDOWN GENERATION
# =============================================================================

def generate_markdown(data: PersonaDocumentData) -> str:
    """Generate markdown version of persona document."""
    total_interviews = sum(p.interview_count for p in data.personas)

    md = f"""# {data.company_name}
## Vianeo Business Model Evaluation

# USER PERSONA ANALYSIS

*{data.project_subtitle}*

---

**Prepared:** {data.prepared_date or format_date(format_str='%B %Y')}
**Validation Status:** {len(data.personas)} personas based on {total_interviews} total interviews
**Evaluation Phase:** Desirability Assessment

VIANEO Framework v2.0

---

## Evidence Base Summary

### Research Overview

{data.research_overview}

### Validation Breakdown

| Persona | Interview Count | Date Range | Validation Status |
|---------|-----------------|------------|-------------------|
"""

    for persona in data.personas:
        status_style = get_validation_style(persona.validation_status)
        md += f"| {persona.first_name} | {persona.interview_count} | {persona.interview_date_range} | {status_style['text']} |\n"

    md += f"""
### Critical Gaps & Next Steps

{data.critical_gaps}

---

"""

    # Add each persona
    for i, persona in enumerate(data.personas, 1):
        status_style = get_validation_style(persona.validation_status)

        md += f"""## PERSONA {i}: {persona.first_name.upper()}

**[{status_style['icon']}] {status_style['text']}**
"""
        if persona.interview_count > 0:
            md += f"*Based on {persona.interview_count} interviews ({persona.interview_date_range})*\n"
        else:
            md += "*No interviews conducted yet*\n"

        md += f"""
### Requester Profile

**First Name:** {persona.first_name}
**Age:** {persona.age}

**Life & Motivations:**
{persona.life_motivations}

**Personality & Values:**
{persona.personality_values}

### Field of Application

**Thinks & Feels:**
{persona.thinks_feels}

**Observes:**
{persona.observes}

**Does:**
{persona.does}

**Others Say:**
{persona.others_say}

### Activities & Challenges

**Tasks (Jobs to be done):**
"""
        for task in persona.tasks[:6]:
            task_text = task[:57] + "..." if len(task) > 60 else task
            md += f"- {task_text}\n"

        md += "\n**Pains (Problems to eliminate):**\n"
        for pain in persona.pains[:6]:
            pain_text = pain[:57] + "..." if len(pain) > 60 else pain
            md += f"- {pain_text}\n"

        md += "\n**Expectations (Desired outcomes):**\n"
        for exp in persona.expectations[:4]:
            exp_text = exp[:57] + "..." if len(exp) > 60 else exp
            md += f"- {exp_text}\n"

        md += f"""
### Current Solutions

{persona.current_solutions_description}

**Current Tools:**
"""
        for tool in persona.current_tools:
            name = tool.get('name', 'Tool')
            limitation = tool.get('limitation', 'Limitation')
            md += f"- **{name}:** {limitation}\n"

        md += "\n---\n\n"

    return md


# =============================================================================
# MAIN FUNCTION
# =============================================================================

def generate_personas(
    input_path: Optional[Path] = None,
    output_path: Optional[Path] = None,
    data: Optional[PersonaDocumentData] = None,
    output_format: str = "both"
) -> Dict[str, Path]:
    """
    Generate Persona document(s).

    Args:
        input_path: Path to input data file (JSON/YAML)
        output_path: Path for output file (without extension)
        data: PersonaDocumentData object (alternative to input_path)
        output_format: "docx", "md", or "both"

    Returns:
        Dict mapping format to output path
    """
    # Load data if not provided
    if data is None:
        if input_path is None:
            raise ValueError("Either input_path or data must be provided")

        input_path = Path(input_path)
        if input_path.suffix in ['.yaml', '.yml']:
            with open(input_path) as f:
                raw_data = yaml.safe_load(f)
        elif input_path.suffix == '.json':
            with open(input_path) as f:
                raw_data = json.load(f)
        else:
            raise ValueError(f"Unsupported input format: {input_path.suffix}")

        # Parse personas
        personas = []
        for p in raw_data.get('personas', []):
            personas.append(PersonaData(**p))

        data = PersonaDocumentData(
            company_name=raw_data.get('company_name', 'Company'),
            project_subtitle=raw_data.get('project_subtitle', ''),
            prepared_date=raw_data.get('prepared_date', ''),
            research_overview=raw_data.get('research_overview', ''),
            critical_gaps=raw_data.get('critical_gaps', ''),
            personas=personas
        )

    # Determine output path
    if output_path is None:
        company_name = safe_filename(data.company_name or "Company")
        date_str = format_date(format_str="%Y%m%d")
        output_path = Path(f"{company_name}_Vianeo_Personas_{date_str}")

    output_path = Path(output_path)
    outputs = {}

    # Generate DOCX
    if output_format in ["docx", "both"] and DOCX_AVAILABLE:
        docx_path = output_path.with_suffix('.docx')
        generator = PersonaDocumentGenerator(data)
        if generator.generate_docx(docx_path):
            outputs['docx'] = docx_path
            print(f"Generated DOCX: {docx_path}")

    # Generate Markdown
    if output_format in ["md", "both"]:
        md_path = output_path.with_suffix('.md')
        md_content = generate_markdown(data)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        outputs['md'] = md_path
        print(f"Generated Markdown: {md_path}")

    return outputs


# =============================================================================
# CLI
# =============================================================================

def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Generate VIANEO Persona documents"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        help='Input data file (JSON or YAML)'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Output file path (without extension)'
    )
    parser.add_argument(
        '--format', '-f',
        choices=['docx', 'md', 'both'],
        default='both',
        help='Output format (default: both)'
    )

    args = parser.parse_args()

    if args.input:
        outputs = generate_personas(
            input_path=args.input,
            output_path=args.output,
            output_format=args.format
        )
        print(f"\nGenerated {len(outputs)} file(s)")
    else:
        print("No input file provided. Use --input to specify data file.")


if __name__ == '__main__':
    main()
