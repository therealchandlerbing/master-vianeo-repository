#!/usr/bin/env python3
"""
VIANEO Executive Sprint Report Generator
=========================================

Generates professional Executive Sprint Report documents for Vianeo Business
Model Evaluation Sprints in both DOCX and Markdown formats.

This report synthesizes findings from a complete Vianeo Sprint (4-6 sessions)
covering all five proof-of-value dimensions.

Usage:
    python generate_executive_sprint_report.py --input sprint_data.yaml --output ExecutiveSprintReport
"""

import argparse
import json
import yaml
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.constants import DocxStyles, ScoreThresholds, VIANEO_DIMENSIONS
from core.utils import format_date, safe_filename, clean_text, load_data_file
from generators.base import BaseDocumentGenerator, is_docx_available, DOCX_AVAILABLE

# Import python-docx components if available
if DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
else:
    # Type hint stub when python-docx not available
    Document = None


# =============================================================================
# DATA MODELS
# =============================================================================

@dataclass
class DimensionDetail:
    """Detailed information for a single dimension assessment."""
    name: str = ""
    weight: str = ""
    score: str = ""
    threshold: str = ""
    status: str = ""
    summary: str = ""
    strengths: List[str] = field(default_factory=list)
    gaps: List[str] = field(default_factory=list)

    @property
    def status_color(self) -> str:
        """Get color for status indicator."""
        if self.status == "PASS":
            return "28A745"  # Green
        return "DC3545"  # Red


@dataclass
class KeyFinding:
    """Key finding for a dimension."""
    dimension: str = ""
    weight: str = ""
    score: str = ""
    status: str = ""
    interpretation: str = ""

    @property
    def status_color(self) -> str:
        """Get color for status indicator."""
        if self.status == "PASS":
            return "28A745"  # Green
        return "DC3545"  # Red


@dataclass
class TargetSegment:
    """Target market segment."""
    segment: str = ""
    characteristics: str = ""


@dataclass
class Persona:
    """Stakeholder persona."""
    name: str = ""
    profile: str = ""
    needs: List[str] = field(default_factory=list)
    validation_required: str = ""


@dataclass
class EcosystemRelationship:
    """Critical ecosystem relationship."""
    relationship: str = ""
    type: str = ""
    criticality: str = ""
    status: str = ""

    @property
    def criticality_color(self) -> str:
        """Get color for criticality level."""
        if self.criticality == "Critical":
            return "DC3545"  # Red
        elif self.criticality == "High":
            return "FFC107"  # Yellow
        return "28A745"  # Green

    @property
    def status_color(self) -> str:
        """Get color for status."""
        if self.status in ["Active", "Funded"]:
            return "28A745"  # Green
        elif self.status in ["Informal", "Not Engaged"]:
            return "FFC107"  # Yellow
        return "6C757D"  # Gray


@dataclass
class Priority:
    """Action priority with detailed breakdown."""
    title: str = ""
    owner: str = ""
    timeline: str = ""
    items_label: str = ""
    items: List[str] = field(default_factory=list)


@dataclass
class MediumTermPriority:
    """Medium-term priority (condensed format)."""
    title: str = ""
    description: str = ""


@dataclass
class RiskMitigation:
    """Risk mitigation strategy."""
    risk: str = ""
    impact: str = ""
    strategy: str = ""

    @property
    def impact_color(self) -> str:
        """Get color for impact level."""
        if self.impact == "Critical":
            return "DC3545"  # Red
        elif self.impact == "High":
            return "FFC107"  # Yellow
        elif self.impact == "Medium":
            return "28A745"  # Green
        return "6C757D"  # Gray


@dataclass
class NextReview:
    """Next review checkpoint information."""
    timing: str = ""
    deliverables: List[str] = field(default_factory=list)
    success_criteria: List[str] = field(default_factory=list)


@dataclass
class ExecutiveSprintReportData:
    """Data structure for VIANEO Executive Sprint Report."""

    # Metadata
    project_name: str = ""
    report_title: str = ""
    report_subtitle: str = ""
    project_tagline: str = ""
    subtitle: str = ""
    principal_investigator: str = ""
    institution: str = ""
    sprint_duration: str = ""
    evaluation_framework: str = "Vianeo Business Model Evaluation System"
    prepared_by: str = "360 Social Impact Studios"
    report_date: str = ""
    author: str = ""
    author_title: str = ""

    # Scores
    overall_vianeo_score: str = ""
    market_maturity_score: str = ""
    status: str = ""

    # Key Findings (for executive summary table)
    key_findings: List[KeyFinding] = field(default_factory=list)

    # Section 1: Executive Summary
    project_overview: List[str] = field(default_factory=list)  # 3 paragraphs
    primary_recommendation_status: str = ""
    primary_recommendation_summary: str = ""
    validation_gaps: List[str] = field(default_factory=list)
    immediate_next_steps: List[str] = field(default_factory=list)

    # Section 2: Business Model Overview
    value_proposition: str = ""
    core_differentiation: List[str] = field(default_factory=list)
    target_segments: List[TargetSegment] = field(default_factory=list)
    revenue_model_type: str = ""
    revenue_model_components: List[str] = field(default_factory=list)
    pricing_warning: str = ""

    # Section 3: Evaluation Results (5 dimensions)
    legitimacy: Optional[DimensionDetail] = None
    desirability: Optional[DimensionDetail] = None
    acceptability: Optional[DimensionDetail] = None
    feasibility: Optional[DimensionDetail] = None
    viability: Optional[DimensionDetail] = None

    # Section 4: Stakeholder Analysis
    personas: List[Persona] = field(default_factory=list)
    ecosystem_relationships: List[EcosystemRelationship] = field(default_factory=list)

    # Section 5: Recommendations
    immediate_priorities: List[Priority] = field(default_factory=list)  # 0-30 days
    short_term_validation: List[Priority] = field(default_factory=list)  # 30-90 days
    medium_term_priorities: List[MediumTermPriority] = field(default_factory=list)  # 90-180 days
    risk_mitigation: List[RiskMitigation] = field(default_factory=list)

    # Section 6: Conclusion
    conclusion: List[str] = field(default_factory=list)  # 4 paragraphs
    next_review: Optional[NextReview] = None


# =============================================================================
# DOCX GENERATION
# =============================================================================

class ExecutiveSprintReportGenerator(BaseDocumentGenerator):
    """Generator for VIANEO Executive Sprint Report documents."""

    def __init__(self, data: ExecutiveSprintReportData):
        super().__init__()
        self.data = data

    def generate_docx(self, output_path: Path) -> bool:
        """Generate professional DOCX executive sprint report."""
        if not is_docx_available():
            print("Error: python-docx not installed")
            return False

        doc = Document()
        self._setup_document(doc)

        # Add sections
        self._add_cover_page(doc)
        self._add_page_break(doc)
        self._add_executive_summary(doc)
        self._add_page_break(doc)
        self._add_business_model_overview(doc)
        self._add_page_break(doc)
        self._add_evaluation_results(doc)
        self._add_page_break(doc)
        self._add_stakeholder_analysis(doc)
        self._add_page_break(doc)
        self._add_recommendations(doc)
        self._add_page_break(doc)
        self._add_conclusion(doc)

        # Save document
        doc.save(str(output_path))
        return True

    def _add_page_break(self, doc: Document) -> None:
        """Add a page break."""
        doc.add_page_break()

    def _add_cover_page(self, doc: Document) -> None:
        """Add cover page with title and metadata table."""
        # Spacing before title
        doc.add_paragraph()
        doc.add_paragraph()

        # Title - Line 1
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.data.report_title)
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("1E3A5F")

        # Title - Line 2
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(self.data.report_subtitle)
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(22)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("1E3A5F")

        # Tagline
        tagline = doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tagline.add_run(self.data.project_tagline)
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(self.styles.LIGHT_GRAY)

        # Assessment type
        assessment = doc.add_paragraph()
        assessment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = assessment.add_run(self.data.subtitle)
        run.font.name = self.styles.FONT_FAMILY
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string("2E5A7F")

        # Add spacing
        doc.add_paragraph()

        # Metadata table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)

        # Populate table
        rows_data = [
            ("Principal Investigator", self.data.principal_investigator),
            ("Institution", self.data.institution),
            ("Sprint Duration", self.data.sprint_duration),
            ("Evaluation Framework", self.data.evaluation_framework),
            ("Prepared By", self.data.prepared_by),
            ("Report Date", self.data.report_date)
        ]

        for i, (label, value) in enumerate(rows_data):
            cells = table.rows[i].cells
            # Label cell (bold)
            cells[0].text = label
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            # Value cell
            cells[1].text = value
            for para in cells[1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

    def _add_executive_summary(self, doc: Document) -> None:
        """Add Section 1: Executive Summary."""
        self._add_styled_heading(doc, "1. Executive Summary", level=1)

        # 1.1 Score Dashboard
        doc.add_heading("1.1 Score Dashboard", level=2)

        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'

        # Set column widths
        for col in table.columns:
            col.width = Inches(2.0)

        # Headers
        headers = ["Overall Vianeo Score", "Market Maturity Score", "Status"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            # Add header shading
            self._add_cell_shading(cell, "1E3A5F")
            self._set_cell_text_color(cell, "FFFFFF")

        # Data
        data_row = [
            self.data.overall_vianeo_score,
            self.data.market_maturity_score,
            self.data.status
        ]
        for i, value in enumerate(data_row):
            cell = table.rows[1].cells[i]
            cell.text = value
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            # Add shading for status column
            if i == 2:
                self._add_cell_shading(cell, "FFC107")

        doc.add_paragraph()

        # 1.2 Project Overview
        doc.add_heading("1.2 Project Overview", level=2)
        for para_text in self.data.project_overview:
            para = doc.add_paragraph()
            run = para.add_run(clean_text(para_text))
            run.font.size = Pt(11)

        # 1.3 Key Findings
        doc.add_heading("1.3 Key Findings", level=2)

        table = doc.add_table(rows=len(self.data.key_findings) + 1, cols=4)
        table.style = 'Table Grid'

        # Set column widths
        widths = [Inches(1.8), Inches(0.8), Inches(1.0), Inches(2.9)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Headers
        headers = ["Dimension", "Score", "Status", "Interpretation"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            self._add_cell_shading(cell, "1E3A5F")
            self._set_cell_text_color(cell, "FFFFFF")

        # Data rows
        for i, finding in enumerate(self.data.key_findings, start=1):
            cells = table.rows[i].cells

            # Dimension
            cells[0].text = f"{finding.dimension} ({finding.weight})"
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            # Score
            cells[1].text = finding.score
            for para in cells[1].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            # Status
            status_text = "✓ PASS" if finding.status == "PASS" else "✗ FAIL"
            cells[2].text = status_text
            for para in cells[2].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor.from_string(finding.status_color)

            # Interpretation
            cells[3].text = finding.interpretation
            for para in cells[3].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

            # Alternating row shading
            if i % 2 == 0:
                for cell in cells:
                    self._add_cell_shading(cell, "E8F4F8")

        doc.add_paragraph()

        # 1.4 Primary Recommendation
        doc.add_heading("1.4 Primary Recommendation", level=2)

        para = doc.add_paragraph()
        run = para.add_run("Status: ")
        run.font.size = Pt(11)
        run = para.add_run(self.data.primary_recommendation_status)
        run.font.size = Pt(11)
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(clean_text(self.data.primary_recommendation_summary))
        run.font.size = Pt(11)

        # Critical validation gaps
        para = doc.add_paragraph()
        run = para.add_run("Critical validation gaps:")
        run.font.size = Pt(11)
        run.bold = True

        for gap in self.data.validation_gaps:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(gap)
            run.font.size = Pt(11)

        # Immediate next steps
        para = doc.add_paragraph()
        run = para.add_run("Immediate next steps (0-90 days):")
        run.font.size = Pt(11)
        run.bold = True

        for step in self.data.immediate_next_steps:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(step)
            run.font.size = Pt(11)

    def _add_business_model_overview(self, doc: Document) -> None:
        """Add Section 2: Business Model Overview."""
        self._add_styled_heading(doc, "2. Business Model Overview", level=1)

        # 2.1 Value Proposition
        doc.add_heading("2.1 Value Proposition", level=2)
        para = doc.add_paragraph()
        run = para.add_run(clean_text(self.data.value_proposition))
        run.font.size = Pt(11)

        para = doc.add_paragraph()
        run = para.add_run("Core differentiation:")
        run.font.size = Pt(11)
        run.bold = True

        for item in self.data.core_differentiation:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(item)
            run.font.size = Pt(11)

        # 2.2 Target Market Segments
        doc.add_heading("2.2 Target Market Segments", level=2)

        table = doc.add_table(rows=len(self.data.target_segments) + 1, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)

        # Headers
        headers = ["Segment", "Key Characteristics"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            self._add_cell_shading(cell, "1E3A5F")
            self._set_cell_text_color(cell, "FFFFFF")

        # Data rows
        for i, segment in enumerate(self.data.target_segments, start=1):
            cells = table.rows[i].cells

            cells[0].text = segment.segment
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            cells[1].text = segment.characteristics
            for para in cells[1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

            # Alternating row shading
            if i % 2 == 0:
                for cell in cells:
                    self._add_cell_shading(cell, "E8F4F8")

        doc.add_paragraph()

        # 2.3 Revenue Model
        doc.add_heading("2.3 Revenue Model", level=2)

        para = doc.add_paragraph()
        run = para.add_run(f"{self.data.revenue_model_type}:")
        run.font.size = Pt(11)
        run.bold = True

        for component in self.data.revenue_model_components:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(component)
            run.font.size = Pt(11)

        # Pricing warning (if present)
        if self.data.pricing_warning:
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.left_indent = Inches(0.25)
            para_format.right_indent = Inches(0.25)

            self._add_cell_shading_paragraph(para, "F8F9FA")

            run = para.add_run("Critical pricing validation needed: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("DC3545")

            run = para.add_run(self.data.pricing_warning)
            run.font.size = Pt(11)

    def _add_evaluation_results(self, doc: Document) -> None:
        """Add Section 3: Evaluation Results."""
        self._add_styled_heading(doc, "3. Evaluation Results by Proof of Value", level=1)

        para = doc.add_paragraph()
        run = para.add_run(
            f"The Vianeo evaluation assessed {self.data.project_name} across five interconnected dimensions, "
            "each weighted according to importance for commercialization success. This section details findings, "
            "evidence, and validation gaps for each proof of value."
        )
        run.font.size = Pt(11)

        # Add each dimension
        dimensions = [
            ("3.1", self.data.legitimacy),
            ("3.2", self.data.desirability),
            ("3.3", self.data.acceptability),
            ("3.4", self.data.feasibility),
            ("3.5", self.data.viability)
        ]

        for section_num, dimension in dimensions:
            if dimension:
                # Page break before each dimension (except first)
                if section_num != "3.1":
                    self._add_page_break(doc)

                self._add_dimension_section(doc, section_num, dimension)

    def _add_dimension_section(self, doc: Document, section_num: str, dim: DimensionDetail) -> None:
        """Add a single dimension section."""
        # Heading
        heading_text = f"{section_num} {dim.name} ({dim.weight}) - Score: {dim.score}"
        doc.add_heading(heading_text, level=2)

        # Status
        para = doc.add_paragraph()
        run = para.add_run("Status: ")
        run.font.size = Pt(11)
        run.bold = True

        status_text = f"✓ {dim.status}" if dim.status == "PASS" else f"✗ {dim.status}"
        run = para.add_run(status_text)
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(dim.status_color)

        run = para.add_run(f" (Threshold: {dim.threshold})")
        run.font.size = Pt(11)

        # Key Findings heading
        doc.add_heading("Key Findings", level=3)

        # Summary
        para = doc.add_paragraph()
        run = para.add_run(clean_text(dim.summary))
        run.font.size = Pt(11)

        # Strengths
        para = doc.add_paragraph()
        run = para.add_run("Strengths:")
        run.font.size = Pt(11)
        run.bold = True

        for strength in dim.strengths:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(strength)
            run.font.size = Pt(11)

        # Gaps
        para = doc.add_paragraph()
        run = para.add_run("Gaps:")
        run.font.size = Pt(11)
        run.bold = True

        for gap in dim.gaps:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(gap)
            run.font.size = Pt(11)

    def _add_stakeholder_analysis(self, doc: Document) -> None:
        """Add Section 4: Stakeholder Analysis."""
        self._add_styled_heading(doc, "4. Stakeholder & Ecosystem Analysis", level=1)

        # 4.1 Priority Personas
        doc.add_heading("4.1 Priority Personas", level=2)

        para = doc.add_paragraph()
        run = para.add_run(
            f"Four primary personas identified for customer discovery validation "
            "(all hypothetical pending interviews):"
        )
        run.font.size = Pt(11)

        # Add each persona
        for i, persona in enumerate(self.data.personas, start=1):
            doc.add_heading(f"{i}. {persona.name}", level=3)

            para = doc.add_paragraph()
            run = para.add_run("Profile: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(persona.profile)
            run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run("Key needs (hypothesized):")
            run.font.size = Pt(11)
            run.bold = True

            for need in persona.needs:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(need)
                run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run("Validation required: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(persona.validation_required)
            run.font.size = Pt(11)

        # 4.2 Critical Ecosystem Relationships
        doc.add_heading("4.2 Critical Ecosystem Relationships", level=2)

        table = doc.add_table(rows=len(self.data.ecosystem_relationships) + 1, cols=4)
        table.style = 'Table Grid'

        # Set column widths
        widths = [Inches(1.8), Inches(1.5), Inches(1.5), Inches(1.7)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Headers
        headers = ["Relationship", "Type", "Criticality", "Status"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            self._add_cell_shading(cell, "1E3A5F")
            self._set_cell_text_color(cell, "FFFFFF")

        # Data rows
        for i, rel in enumerate(self.data.ecosystem_relationships, start=1):
            cells = table.rows[i].cells

            cells[0].text = rel.relationship
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            cells[1].text = rel.type
            for para in cells[1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

            # Criticality with color
            cells[2].text = rel.criticality
            for para in cells[2].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor.from_string(rel.criticality_color)

            # Status with color
            cells[3].text = rel.status
            for para in cells[3].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor.from_string(rel.status_color)

            # Alternating row shading
            if i % 2 == 0:
                for cell in cells:
                    self._add_cell_shading(cell, "E8F4F8")

    def _add_recommendations(self, doc: Document) -> None:
        """Add Section 5: Recommendations."""
        self._add_styled_heading(doc, "5. Recommendations & Next Steps", level=1)

        # 5.1 Immediate Priorities
        doc.add_heading("5.1 Immediate Priorities (0-30 Days)", level=2)

        for i, priority in enumerate(self.data.immediate_priorities, start=1):
            doc.add_heading(f"{i}. {priority.title}", level=3)

            para = doc.add_paragraph()
            run = para.add_run("Owner: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(priority.owner)
            run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run("Timeline: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(priority.timeline)
            run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run(f"{priority.items_label}:")
            run.font.size = Pt(11)
            run.bold = True

            for item in priority.items:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(item)
                run.font.size = Pt(11)

        # 5.2 Short-Term Validation
        doc.add_heading("5.2 Short-Term Validation (30-90 Days)", level=2)

        start_num = len(self.data.immediate_priorities) + 1
        for i, priority in enumerate(self.data.short_term_validation, start=start_num):
            doc.add_heading(f"{i}. {priority.title}", level=3)

            para = doc.add_paragraph()
            run = para.add_run("Owner: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(priority.owner)
            run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run("Timeline: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(priority.timeline)
            run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run(f"{priority.items_label}:")
            run.font.size = Pt(11)
            run.bold = True

            for item in priority.items:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(item)
                run.font.size = Pt(11)

        # Page break before medium-term
        self._add_page_break(doc)

        # 5.3 Medium-Term Priorities
        doc.add_heading("5.3 Medium-Term Priorities (90-180 Days)", level=2)

        start_num = len(self.data.immediate_priorities) + len(self.data.short_term_validation) + 1
        for i, priority in enumerate(self.data.medium_term_priorities, start=start_num):
            doc.add_heading(f"{i}. {priority.title}", level=3)

            para = doc.add_paragraph()
            run = para.add_run(clean_text(priority.description))
            run.font.size = Pt(11)

        # 5.4 Risk Mitigation
        doc.add_heading("5.4 Risk Mitigation Strategies", level=2)

        table = doc.add_table(rows=len(self.data.risk_mitigation) + 1, cols=3)
        table.style = 'Table Grid'

        # Set column widths
        widths = [Inches(2.0), Inches(1.0), Inches(3.5)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Headers
        headers = ["Risk", "Impact", "Mitigation Strategy"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            self._add_cell_shading(cell, "1E3A5F")
            self._set_cell_text_color(cell, "FFFFFF")

        # Data rows
        for i, risk in enumerate(self.data.risk_mitigation, start=1):
            cells = table.rows[i].cells

            cells[0].text = risk.risk
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            # Impact with color
            cells[1].text = risk.impact
            for para in cells[1].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor.from_string(risk.impact_color)

            cells[2].text = risk.strategy
            for para in cells[2].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

            # Alternating row shading
            if i % 2 == 0:
                for cell in cells:
                    self._add_cell_shading(cell, "E8F4F8")

    def _add_conclusion(self, doc: Document) -> None:
        """Add Section 6: Conclusion."""
        self._add_styled_heading(doc, "6. Conclusion", level=1)

        # Conclusion paragraphs
        for i, para_text in enumerate(self.data.conclusion):
            para = doc.add_paragraph()

            # Special formatting for paragraphs 3 and 4
            if i == 2 and "path forward is clear" in para_text.lower():
                run = para.add_run("The path forward is clear: ")
                run.font.size = Pt(11)
                run.bold = True
                run = para.add_run(para_text.replace("The path forward is clear: ", ""))
                run.font.size = Pt(11)
            elif i == 3 and "CONDITIONAL PROCEED" in para_text:
                run = para.add_run("The assessment recommends ")
                run.font.size = Pt(11)
                run = para.add_run("CONDITIONAL PROCEED")
                run.font.size = Pt(11)
                run.bold = True
                remaining = para_text.replace("The assessment recommends CONDITIONAL PROCEED", "")
                run = para.add_run(remaining)
                run.font.size = Pt(11)
            else:
                run = para.add_run(clean_text(para_text))
                run.font.size = Pt(11)

        # Next Review Checkpoint
        if self.data.next_review:
            doc.add_heading("6.1 Next Review Checkpoint", level=2)

            para = doc.add_paragraph()
            run = para.add_run("Timing: ")
            run.font.size = Pt(11)
            run.bold = True
            run = para.add_run(self.data.next_review.timing)
            run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run("Expected Deliverables:")
            run.font.size = Pt(11)
            run.bold = True

            for deliverable in self.data.next_review.deliverables:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(deliverable)
                run.font.size = Pt(11)

            para = doc.add_paragraph()
            run = para.add_run("Success Criteria for Series A Readiness:")
            run.font.size = Pt(11)
            run.bold = True

            for criterion in self.data.next_review.success_criteria:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(criterion)
                run.font.size = Pt(11)

        # Footer
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("— End of Report —")
        run.font.size = Pt(11)
        run.italic = True

        doc.add_paragraph()

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Prepared by {self.data.prepared_by}")
        run.font.size = Pt(11)
        run.bold = True

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"{self.data.author}, {self.data.author_title}")
        run.font.size = Pt(11)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("Using Vianeo Business Model Evaluation Framework")
        run.font.size = Pt(11)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.data.report_date)
        run.font.size = Pt(11)

    def _add_cell_shading(self, cell, color: str) -> None:
        """Add shading to a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_text_color(self, cell, color: str) -> None:
        """Set text color for all runs in a cell."""
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(color)

    def _add_cell_shading_paragraph(self, para, color: str) -> None:
        """Add shading to a paragraph (for callout boxes)."""
        # Note: This is a simplified version - proper implementation would use borders
        pass


# =============================================================================
# MARKDOWN GENERATION
# =============================================================================

def generate_markdown(data: ExecutiveSprintReportData) -> str:
    """Generate markdown version of executive sprint report."""

    md = f"""# {data.report_title}
## {data.report_subtitle}

*{data.project_tagline}*

**{data.subtitle}**

---

| Field | Value |
|-------|-------|
| **Principal Investigator** | {data.principal_investigator} |
| **Institution** | {data.institution} |
| **Sprint Duration** | {data.sprint_duration} |
| **Evaluation Framework** | {data.evaluation_framework} |
| **Prepared By** | {data.prepared_by} |
| **Report Date** | {data.report_date} |

---

## 1. Executive Summary

### 1.1 Score Dashboard

| Overall Vianeo Score | Market Maturity Score | Status |
|---------------------|----------------------|--------|
| **{data.overall_vianeo_score}** | **{data.market_maturity_score}** | **{data.status}** |

### 1.2 Project Overview

"""

    for para in data.project_overview:
        md += f"{para}\n\n"

    md += "### 1.3 Key Findings\n\n"
    md += "| Dimension | Score | Status | Interpretation |\n"
    md += "|-----------|-------|--------|----------------|\n"

    for finding in data.key_findings:
        status_mark = "✓ PASS" if finding.status == "PASS" else "✗ FAIL"
        md += f"| **{finding.dimension} ({finding.weight})** | {finding.score} | {status_mark} | {finding.interpretation} |\n"

    md += f"""

### 1.4 Primary Recommendation

**Status:** {data.primary_recommendation_status}

{data.primary_recommendation_summary}

**Critical validation gaps:**

"""

    for gap in data.validation_gaps:
        md += f"- {gap}\n"

    md += "\n**Immediate next steps (0-90 days):**\n\n"

    for step in data.immediate_next_steps:
        md += f"- {step}\n"

    md += """

---

## 2. Business Model Overview

### 2.1 Value Proposition

"""

    md += f"{data.value_proposition}\n\n"
    md += "**Core differentiation:**\n\n"

    for item in data.core_differentiation:
        md += f"- {item}\n"

    md += "\n### 2.2 Target Market Segments\n\n"
    md += "| Segment | Key Characteristics |\n"
    md += "|---------|---------------------|\n"

    for seg in data.target_segments:
        md += f"| **{seg.segment}** | {seg.characteristics} |\n"

    md += f"\n### 2.3 Revenue Model\n\n**{data.revenue_model_type}:**\n\n"

    for comp in data.revenue_model_components:
        md += f"- {comp}\n"

    if data.pricing_warning:
        md += f"\n> **Critical pricing validation needed:** {data.pricing_warning}\n"

    md += "\n---\n\n## 3. Evaluation Results by Proof of Value\n\n"
    md += f"The Vianeo evaluation assessed {data.project_name} across five interconnected dimensions, "
    md += "each weighted according to importance for commercialization success. This section details findings, "
    md += "evidence, and validation gaps for each proof of value.\n\n"

    # Add dimensions
    dimensions = [
        ("3.1", data.legitimacy),
        ("3.2", data.desirability),
        ("3.3", data.acceptability),
        ("3.4", data.feasibility),
        ("3.5", data.viability)
    ]

    for section_num, dim in dimensions:
        if dim:
            status_mark = "✓ PASS" if dim.status == "PASS" else "✗ FAIL"
            md += f"### {section_num} {dim.name} ({dim.weight}) - Score: {dim.score}\n\n"
            md += f"**Status:** {status_mark} (Threshold: {dim.threshold})\n\n"
            md += "#### Key Findings\n\n"
            md += f"{dim.summary}\n\n"
            md += "**Strengths:**\n\n"
            for strength in dim.strengths:
                md += f"- {strength}\n"
            md += "\n**Gaps:**\n\n"
            for gap in dim.gaps:
                md += f"- {gap}\n"
            md += "\n"

    md += "---\n\n## 4. Stakeholder & Ecosystem Analysis\n\n"
    md += "### 4.1 Priority Personas\n\n"
    md += "Four primary personas identified for customer discovery validation (all hypothetical pending interviews):\n\n"

    for i, persona in enumerate(data.personas, start=1):
        md += f"#### {i}. {persona.name}\n\n"
        md += f"**Profile:** {persona.profile}\n\n"
        md += "**Key needs (hypothesized):**\n\n"
        for need in persona.needs:
            md += f"- {need}\n"
        md += f"\n**Validation required:** {persona.validation_required}\n\n"

    md += "### 4.2 Critical Ecosystem Relationships\n\n"
    md += "| Relationship | Type | Criticality | Status |\n"
    md += "|--------------|------|-------------|--------|\n"

    for rel in data.ecosystem_relationships:
        md += f"| **{rel.relationship}** | {rel.type} | {rel.criticality} | {rel.status} |\n"

    md += "\n---\n\n## 5. Recommendations & Next Steps\n\n"
    md += "### 5.1 Immediate Priorities (0-30 Days)\n\n"

    for i, priority in enumerate(data.immediate_priorities, start=1):
        md += f"#### {i}. {priority.title}\n\n"
        md += f"**Owner:** {priority.owner}\n\n"
        md += f"**Timeline:** {priority.timeline}\n\n"
        md += f"**{priority.items_label}:**\n\n"
        for item in priority.items:
            md += f"- {item}\n"
        md += "\n"

    md += "### 5.2 Short-Term Validation (30-90 Days)\n\n"

    start_num = len(data.immediate_priorities) + 1
    for i, priority in enumerate(data.short_term_validation, start=start_num):
        md += f"#### {i}. {priority.title}\n\n"
        md += f"**Owner:** {priority.owner}\n\n"
        md += f"**Timeline:** {priority.timeline}\n\n"
        md += f"**{priority.items_label}:**\n\n"
        for item in priority.items:
            md += f"- {item}\n"
        md += "\n"

    md += "### 5.3 Medium-Term Priorities (90-180 Days)\n\n"

    start_num = len(data.immediate_priorities) + len(data.short_term_validation) + 1
    for i, priority in enumerate(data.medium_term_priorities, start=start_num):
        md += f"#### {i}. {priority.title}\n\n"
        md += f"{priority.description}\n\n"

    md += "### 5.4 Risk Mitigation Strategies\n\n"
    md += "| Risk | Impact | Mitigation Strategy |\n"
    md += "|------|--------|---------------------|\n"

    for risk in data.risk_mitigation:
        md += f"| **{risk.risk}** | {risk.impact} | {risk.strategy} |\n"

    md += "\n---\n\n## 6. Conclusion\n\n"

    for para in data.conclusion:
        md += f"{para}\n\n"

    if data.next_review:
        md += "### 6.1 Next Review Checkpoint\n\n"
        md += f"**Timing:** {data.next_review.timing}\n\n"
        md += "**Expected Deliverables:**\n\n"
        for deliverable in data.next_review.deliverables:
            md += f"- {deliverable}\n"
        md += "\n**Success Criteria for Series A Readiness:**\n\n"
        for criterion in data.next_review.success_criteria:
            md += f"- {criterion}\n"

    md += f"""

---

*— End of Report —*

**Prepared by {data.prepared_by}**
{data.author}, {data.author_title}
Using Vianeo Business Model Evaluation Framework
{data.report_date}
"""

    return md


# =============================================================================
# MAIN FUNCTION
# =============================================================================

def generate_executive_sprint_report(
    input_path: Optional[Path] = None,
    output_path: Optional[Path] = None,
    data: Optional[ExecutiveSprintReportData] = None,
    output_format: str = "both"
) -> Dict[str, Path]:
    """
    Generate Executive Sprint Report document(s).

    Args:
        input_path: Path to input data file (JSON/YAML)
        output_path: Path for output file (without extension)
        data: ExecutiveSprintReportData object (alternative to input_path)
        output_format: "docx", "md", or "both"

    Returns:
        Dict mapping format to output path
    """
    # Load data if not provided
    if data is None:
        if input_path is None:
            raise ValueError("Either input_path or data must be provided")

        raw_data = load_data_file(input_path)

        # Parse nested objects
        key_findings = [KeyFinding(**kf) for kf in raw_data.get('key_findings', [])]
        target_segments = [TargetSegment(**ts) for ts in raw_data.get('target_segments', [])]
        personas = [Persona(**p) for p in raw_data.get('personas', [])]
        ecosystem_relationships = [EcosystemRelationship(**er) for er in raw_data.get('ecosystem_relationships', [])]
        immediate_priorities = [Priority(**ip) for ip in raw_data.get('immediate_priorities', [])]
        short_term_validation = [Priority(**stv) for stv in raw_data.get('short_term_validation', [])]
        medium_term_priorities = [MediumTermPriority(**mtp) for mtp in raw_data.get('medium_term_priorities', [])]
        risk_mitigation = [RiskMitigation(**rm) for rm in raw_data.get('risk_mitigation', [])]

        # Parse dimension details
        legitimacy = DimensionDetail(**raw_data['legitimacy']) if 'legitimacy' in raw_data else None
        desirability = DimensionDetail(**raw_data['desirability']) if 'desirability' in raw_data else None
        acceptability = DimensionDetail(**raw_data['acceptability']) if 'acceptability' in raw_data else None
        feasibility = DimensionDetail(**raw_data['feasibility']) if 'feasibility' in raw_data else None
        viability = DimensionDetail(**raw_data['viability']) if 'viability' in raw_data else None

        # Parse next review
        next_review = NextReview(**raw_data['next_review']) if 'next_review' in raw_data else None

        data = ExecutiveSprintReportData(
            project_name=raw_data.get('project_name', ''),
            report_title=raw_data.get('report_title', ''),
            report_subtitle=raw_data.get('report_subtitle', ''),
            project_tagline=raw_data.get('project_tagline', ''),
            subtitle=raw_data.get('subtitle', ''),
            principal_investigator=raw_data.get('principal_investigator', ''),
            institution=raw_data.get('institution', ''),
            sprint_duration=raw_data.get('sprint_duration', ''),
            evaluation_framework=raw_data.get('evaluation_framework', 'Vianeo Business Model Evaluation System'),
            prepared_by=raw_data.get('prepared_by', '360 Social Impact Studios'),
            report_date=raw_data.get('report_date', ''),
            author=raw_data.get('author', ''),
            author_title=raw_data.get('author_title', ''),
            overall_vianeo_score=raw_data.get('overall_vianeo_score', ''),
            market_maturity_score=raw_data.get('market_maturity_score', ''),
            status=raw_data.get('status', ''),
            key_findings=key_findings,
            project_overview=raw_data.get('project_overview', []),
            primary_recommendation_status=raw_data.get('primary_recommendation_status', ''),
            primary_recommendation_summary=raw_data.get('primary_recommendation_summary', ''),
            validation_gaps=raw_data.get('validation_gaps', []),
            immediate_next_steps=raw_data.get('immediate_next_steps', []),
            value_proposition=raw_data.get('value_proposition', ''),
            core_differentiation=raw_data.get('core_differentiation', []),
            target_segments=target_segments,
            revenue_model_type=raw_data.get('revenue_model_type', ''),
            revenue_model_components=raw_data.get('revenue_model_components', []),
            pricing_warning=raw_data.get('pricing_warning', ''),
            legitimacy=legitimacy,
            desirability=desirability,
            acceptability=acceptability,
            feasibility=feasibility,
            viability=viability,
            personas=personas,
            ecosystem_relationships=ecosystem_relationships,
            immediate_priorities=immediate_priorities,
            short_term_validation=short_term_validation,
            medium_term_priorities=medium_term_priorities,
            risk_mitigation=risk_mitigation,
            conclusion=raw_data.get('conclusion', []),
            next_review=next_review
        )

    # Determine output path
    if output_path is None:
        project_name = safe_filename(data.project_name or "Project")
        today = datetime.now().strftime("%Y%m%d")
        output_path = Path(f"{project_name}_Vianeo_Sprint_Executive_Report_{today}")

    output_path = Path(output_path)
    outputs = {}

    # Generate Markdown first (always available)
    if output_format in ["md", "both"]:
        md_path = output_path.with_suffix('.md')
        md_content = generate_markdown(data)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        outputs['md'] = md_path
        print(f"Generated Markdown: {md_path}")

    # Generate DOCX
    if output_format in ["docx", "both"] and is_docx_available():
        docx_path = output_path.with_suffix('.docx')
        generator = ExecutiveSprintReportGenerator(data)
        if generator.generate_docx(docx_path):
            outputs['docx'] = docx_path
            print(f"Generated DOCX: {docx_path}")

    return outputs


# =============================================================================
# CLI
# =============================================================================

def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Generate VIANEO Executive Sprint Report documents"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        help='Input data file (JSON or YAML)'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Output file path (without extension)'
    )
    parser.add_argument(
        '--format', '-f',
        choices=['docx', 'md', 'both'],
        default='both',
        help='Output format (default: both)'
    )

    args = parser.parse_args()

    if args.input:
        outputs = generate_executive_sprint_report(
            input_path=args.input,
            output_path=args.output,
            output_format=args.format
        )
        print(f"\nGenerated {len(outputs)} file(s)")
    else:
        print("No input file provided. Use --input to specify data file.")


if __name__ == '__main__':
    main()
